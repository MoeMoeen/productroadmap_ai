# brain/services/document_processor.py

import os
import time
from typing import List, Dict, Any
import mimetypes

# Document processing libraries
import pdfplumber
from docx import Document as DocxDocument
from openpyxl import load_workbook

# Local imports
from brain.models.documents import ParsedDocument, DocumentMetadata, ValidationResult


class DocumentProcessingError(Exception):
    """Custom exception for document processing errors"""
    pass


class DocumentProcessor:
    """
    Core document processing service for the Perception Layer.
    Handles PDF, DOCX, and XLSX files with comprehensive error handling.
    """
    
    SUPPORTED_MIME_TYPES = {
        'application/pdf': 'pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
        'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': 'xlsx',
        'application/msword': 'doc',  # Limited support
        'text/plain': 'txt',
    }
    
    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
    MAX_PAGES_PDF = 500
    MAX_SHEETS_XLSX = 100
    
    def __init__(self):
        self.stats = {
            'files_processed': 0,
            'total_pages': 0,
            'total_tables': 0,
            'processing_time': 0.0,
            'errors': []
        }
    
    def process_files(self, file_paths: List[str]) -> List[ParsedDocument]:
        """
        Process multiple files with comprehensive error handling.
        
        Args:
            file_paths: List of absolute file paths to process
            
        Returns:
            List of ParsedDocument objects with metadata and validation results
        """
        start_time = time.perf_counter()
        processed_documents = []
        
        for file_path in file_paths:
            try:
                document = self._process_single_file(file_path)
                processed_documents.append(document)
                self.stats['files_processed'] += 1
                
            except DocumentProcessingError as e:
                # Create failed document with error details
                failed_doc = self._create_failed_document(file_path, str(e))
                processed_documents.append(failed_doc)
                self.stats['errors'].append(f"{file_path}: {e}")
                
            except Exception as e:
                # Unexpected error - log and create failed document
                error_msg = f"Unexpected error processing {file_path}: {e}"
                failed_doc = self._create_failed_document(file_path, error_msg)
                processed_documents.append(failed_doc)
                self.stats['errors'].append(error_msg)
        
        self.stats['processing_time'] = time.perf_counter() - start_time
        return processed_documents
    
    def _process_single_file(self, file_path: str) -> ParsedDocument:
        """Process a single file based on its type."""
        # Validate file exists and get basic info
        if not os.path.exists(file_path):
            raise DocumentProcessingError(f"File not found: {file_path}")
        
        file_size = os.path.getsize(file_path)
        if file_size > self.MAX_FILE_SIZE:
            raise DocumentProcessingError(f"File too large: {file_size} bytes (max: {self.MAX_FILE_SIZE})")
        
        # Determine file type
        mime_type, _ = mimetypes.guess_type(file_path)
        file_type = self.SUPPORTED_MIME_TYPES.get(mime_type or "")
        
        if not file_type:
            raise DocumentProcessingError(f"Unsupported file type: {mime_type}")
        
        # Process based on file type
        if file_type == 'pdf':
            return self._process_pdf(file_path)
        elif file_type == 'docx':
            return self._process_docx(file_path)
        elif file_type == 'xlsx':
            return self._process_xlsx(file_path)
        elif file_type == 'txt':
            return self._process_txt(file_path)
        else:
            raise DocumentProcessingError(f"Handler not implemented for file type: {file_type}")
    
    def _process_pdf(self, file_path: str) -> ParsedDocument:
        """Extract text, tables, and metadata from PDF files."""
        start_time = time.perf_counter()
        content_parts = []
        tables = []
        page_count = 0
        
        try:
            with pdfplumber.open(file_path) as pdf:
                page_count = len(pdf.pages)
                
                if page_count > self.MAX_PAGES_PDF:
                    raise DocumentProcessingError(f"PDF too large: {page_count} pages (max: {self.MAX_PAGES_PDF})")
                
                for page_num, page in enumerate(pdf.pages, 1):
                    # Extract text
                    page_text = page.extract_text()
                    if page_text:
                        content_parts.append(f"--- Page {page_num} ---\n{page_text.strip()}")
                    
                    # Extract tables
                    page_tables = page.extract_tables()
                    for table_idx, table in enumerate(page_tables):
                        if table and len(table) > 1:  # Skip empty or single-row tables
                            tables.append({
                                'page': page_num,
                                'table_index': table_idx,
                                'headers': table[0] if table else [],
                                'rows': table[1:] if len(table) > 1 else [],
                                'row_count': len(table) - 1
                            })
                
        except Exception as e:
            raise DocumentProcessingError(f"PDF processing failed: {e}")
        
        processing_time = time.perf_counter() - start_time
        content = "\n\n".join(content_parts)
        
        # Create metadata
        metadata = DocumentMetadata(
            file_path=file_path,
            file_size=os.path.getsize(file_path),
            file_type="pdf",
            page_count=page_count,
            table_count=len(tables),
            processing_time_ms=int(processing_time * 1000),
            extracted_text_length=len(content)
        )
        
        # Validate extraction quality
        validation = self._validate_extraction(content, tables, metadata)
        
        self.stats['total_pages'] += page_count
        self.stats['total_tables'] += len(tables)
        
        return ParsedDocument(
            file_path=file_path,
            file_type="pdf",
            content=content,
            tables=tables,
            metadata=metadata,
            validation_result=validation
        )
    
    def _process_docx(self, file_path: str) -> ParsedDocument:
        """Extract structured content from Word documents."""
        start_time = time.perf_counter()
        content_parts = []
        tables = []
        
        try:
            doc = DocxDocument(file_path)
            
            # Extract paragraphs and maintain structure
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    # Preserve heading styles if available
                    style = para.style.name if para.style else "Normal"
                    if style and "Heading" in style:
                        content_parts.append(f"\n## {text}\n")
                    else:
                        content_parts.append(text)
            
            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                
                if table_data:
                    tables.append({
                        'table_index': table_idx,
                        'headers': table_data[0] if table_data else [],
                        'rows': table_data[1:] if len(table_data) > 1 else [],
                        'row_count': len(table_data) - 1
                    })
                    
        except Exception as e:
            raise DocumentProcessingError(f"DOCX processing failed: {e}")
        
        processing_time = time.perf_counter() - start_time
        content = "\n".join(content_parts)
        
        # Create metadata
        metadata = DocumentMetadata(
            file_path=file_path,
            file_size=os.path.getsize(file_path),
            file_type="docx",
            page_count=1,  # DOCX doesn't have explicit pages
            table_count=len(tables),
            processing_time_ms=int(processing_time * 1000),
            extracted_text_length=len(content)
        )
        
        # Validate extraction
        validation = self._validate_extraction(content, tables, metadata)
        
        self.stats['total_pages'] += 1
        self.stats['total_tables'] += len(tables)
        
        return ParsedDocument(
            file_path=file_path,
            file_type="docx",
            content=content,
            tables=tables,
            metadata=metadata,
            validation_result=validation
        )
    
    def _process_xlsx(self, file_path: str) -> ParsedDocument:
        """Parse spreadsheet data with sheet detection."""
        start_time = time.perf_counter()
        content_parts = []
        tables = []
        
        try:
            workbook = load_workbook(file_path, data_only=True)
            sheet_count = len(workbook.worksheets)
            
            if sheet_count > self.MAX_SHEETS_XLSX:
                raise DocumentProcessingError(f"Excel file too large: {sheet_count} sheets (max: {self.MAX_SHEETS_XLSX})")
            
            for sheet_idx, worksheet in enumerate(workbook.worksheets):
                sheet_name = worksheet.title
                content_parts.append(f"\n--- Sheet: {sheet_name} ---")
                
                # Get all data from sheet
                sheet_data = []
                for row in worksheet.iter_rows(values_only=True):
                    # Filter out completely empty rows
                    if any(cell is not None and str(cell).strip() for cell in row):
                        sheet_data.append([str(cell) if cell is not None else "" for cell in row])
                
                if sheet_data:
                    # Create table representation
                    tables.append({
                        'sheet_name': sheet_name,
                        'sheet_index': sheet_idx,
                        'headers': sheet_data[0] if sheet_data else [],
                        'rows': sheet_data[1:] if len(sheet_data) > 1 else [],
                        'row_count': len(sheet_data) - 1,
                        'col_count': len(sheet_data[0]) if sheet_data else 0
                    })
                    
                    # Add summary to content
                    row_count = len(sheet_data)
                    col_count = len(sheet_data[0]) if sheet_data else 0
                    content_parts.append(f"Data: {row_count} rows × {col_count} columns")
                    
                    # Add first few rows as sample
                    if len(sheet_data) > 0:
                        content_parts.append("Sample data:")
                        for i, row in enumerate(sheet_data[:5]):  # First 5 rows
                            content_parts.append(f"  Row {i+1}: {', '.join(row[:10])}")  # First 10 columns
                            
        except Exception as e:
            raise DocumentProcessingError(f"XLSX processing failed: {e}")
        
        processing_time = time.perf_counter() - start_time
        content = "\n".join(content_parts)
        
        # Create metadata
        metadata = DocumentMetadata(
            file_path=file_path,
            file_size=os.path.getsize(file_path),
            file_type="xlsx",
            page_count=sheet_count,
            table_count=len(tables),
            processing_time_ms=int(processing_time * 1000),
            extracted_text_length=len(content)
        )
        
        # Validate extraction
        validation = self._validate_extraction(content, tables, metadata)
        
        self.stats['total_pages'] += sheet_count
        self.stats['total_tables'] += len(tables)
        
        return ParsedDocument(
            file_path=file_path,
            file_type="xlsx",
            content=content,
            tables=tables,
            metadata=metadata,
            validation_result=validation
        )
    
    def _process_txt(self, file_path: str) -> ParsedDocument:
        """Process plain text files."""
        start_time = time.perf_counter()
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    content = f.read()
            except Exception as e:
                raise DocumentProcessingError(f"Text file encoding error: {e}")
        except Exception as e:
            raise DocumentProcessingError(f"Text file processing failed: {e}")
        
        processing_time = time.perf_counter() - start_time
        
        # Create metadata
        metadata = DocumentMetadata(
            file_path=file_path,
            file_size=os.path.getsize(file_path),
            file_type="txt",
            page_count=1,
            table_count=0,
            processing_time_ms=int(processing_time * 1000),
            extracted_text_length=len(content)
        )
        
        # Simple validation for text files
        validation = ValidationResult(
            is_valid=True,
            quality_score=1.0 if content.strip() else 0.0,
            errors=[],
            warnings=[] if content.strip() else ["Text file appears to be empty"],
            details={"encoding": "utf-8", "line_count": content.count('\n') + 1}
        )
        
        return ParsedDocument(
            file_path=file_path,
            file_type="txt",
            content=content,
            tables=[],
            metadata=metadata,
            validation_result=validation
        )
    
    def _validate_extraction(self, content: str, tables: List[Dict], metadata: DocumentMetadata) -> ValidationResult:
        """Validate the quality of extraction results."""
        errors = []
        warnings = []
        quality_score = 1.0
        
        # Content validation
        if not content.strip():
            errors.append("No text content extracted")
            quality_score *= 0.0
        elif len(content.strip()) < 50:
            warnings.append("Very little text content extracted")
            quality_score *= 0.7
        
        # File size vs content ratio check
        content_ratio = len(content) / max(metadata.file_size, 1)
        if content_ratio < 0.001:  # Less than 0.1% conversion rate
            warnings.append("Low text extraction ratio - file may contain mostly images or formatting")
            quality_score *= 0.8
        
        # Processing time check
        if metadata.processing_time_ms > 30000:  # 30 seconds
            warnings.append("Processing took longer than expected")
        
        # Table validation
        if metadata.file_type in ['xlsx', 'pdf'] and not tables:
            warnings.append("No tables found - this may be expected or indicate parsing issues")
        
        is_valid = len(errors) == 0
        
        return ValidationResult(
            is_valid=is_valid,
            quality_score=quality_score,
            errors=errors,
            warnings=warnings,
            details={
                "content_length": len(content),
                "content_ratio": content_ratio,
                "processing_time_ms": metadata.processing_time_ms,
                "table_count": len(tables)
            }
        )
    
    def _create_failed_document(self, file_path: str, error_message: str) -> ParsedDocument:
        """Create a ParsedDocument for failed processing."""
        metadata = DocumentMetadata(
            file_path=file_path,
            file_size=os.path.getsize(file_path) if os.path.exists(file_path) else 0,
            file_type="unknown",
            page_count=0,
            table_count=0,
            processing_time_ms=0,
            extracted_text_length=0
        )
        
        validation = ValidationResult(
            is_valid=False,
            quality_score=0.0,
            errors=[error_message],
            warnings=[],
            details={"processing_failed": True}
        )
        
        return ParsedDocument(
            file_path=file_path,
            file_type="failed",
            content="",
            tables=[],
            metadata=metadata,
            validation_result=validation
        )
    
    def get_processing_stats(self) -> Dict[str, Any]:
        """Get processing statistics for monitoring and debugging."""
        return {
            **self.stats,
            'success_rate': (self.stats['files_processed'] - len(self.stats['errors'])) / max(self.stats['files_processed'], 1),
            'avg_processing_time': self.stats['processing_time'] / max(self.stats['files_processed'], 1) if self.stats['files_processed'] > 0 else 0
        }
