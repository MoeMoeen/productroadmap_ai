#!/usr/bin/env python3
"""
Create comprehensive test documents for production testing
Generates PDF, DOCX, and XLSX files with realistic product roadmap content
"""

import os
from docx import Document
from openpyxl import Workbook
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors

def create_pdf_document(filename):
    """Create a comprehensive PDF document with product roadmap content"""
    doc = SimpleDocTemplate(filename, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    # Title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        spaceAfter=30,
        alignment=1  # Center alignment
    )
    story.append(Paragraph("Product Roadmap 2024-2025", title_style))
    story.append(Spacer(1, 12))
    
    # Executive Summary
    story.append(Paragraph("Executive Summary", styles['Heading2']))
    story.append(Paragraph(
        "This document outlines the strategic product roadmap for our AI-powered document "
        "processing platform. The roadmap focuses on three key areas: enhanced AI capabilities, "
        "improved user experience, and scalable infrastructure. Key initiatives include implementing "
        "hybrid LLM processing, developing real-time collaboration features, and expanding "
        "multi-tenant support. Expected outcomes include 40% improvement in processing accuracy, "
        "60% reduction in processing time, and support for 10x more concurrent users.",
        styles['Normal']
    ))
    story.append(Spacer(1, 12))
    
    # Key Features
    story.append(Paragraph("Key Features & Initiatives", styles['Heading2']))
    features = [
        ["Feature", "Priority", "Timeline", "Impact Score"],
        ["Hybrid LLM Processing", "High", "Q1 2024", "9/10"],
        ["Real-time Collaboration", "Medium", "Q2 2024", "7/10"],
        ["Advanced Analytics Dashboard", "High", "Q1 2024", "8/10"],
        ["Multi-format Support", "Medium", "Q3 2024", "6/10"],
        ["API Rate Limiting", "Low", "Q4 2024", "5/10"]
    ]
    
    table = Table(features, colWidths=[2.5*inch, 1*inch, 1*inch, 1*inch])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    story.append(table)
    story.append(Spacer(1, 12))
    
    # Technical Requirements
    story.append(Paragraph("Technical Requirements", styles['Heading2']))
    story.append(Paragraph(
        "• Python 3.10+ with Django 5.x framework\n"
        "• LangGraph for workflow orchestration\n"
        "• Anthropic Claude and OpenAI GPT-4 integration\n"
        "• PostgreSQL for data persistence\n"
        "• Redis for caching and session management\n"
        "• Docker containerization for deployment\n"
        "• Kubernetes for orchestration and scaling",
        styles['Normal']
    ))
    
    doc.build(story)
    print(f"Created PDF document: {filename}")

def create_docx_document(filename):
    """Create a comprehensive DOCX document with user stories and requirements"""
    doc = Document()
    
    # Title
    title = doc.add_heading('User Stories & Requirements Document', 0)
    
    # Introduction
    doc.add_heading('Product Vision', level=1)
    intro = doc.add_paragraph(
        "Our AI-powered document processing platform revolutionizes how organizations "
        "handle unstructured data. By combining traditional parsing methods with "
        "advanced LLM capabilities, we provide accurate, fast, and scalable document "
        "processing solutions."
    )
    
    # User Stories
    doc.add_heading('User Stories', level=1)
    
    stories = [
        {
            'title': 'Document Upload and Processing',
            'story': 'As a product manager, I want to upload multiple documents (PDF, DOCX, XLSX) '
                    'so that I can extract insights for roadmap planning.',
            'acceptance': [
                'User can upload multiple files simultaneously',
                'System supports PDF, DOCX, and XLSX formats',
                'Processing status is visible in real-time',
                'Extracted content is structured and searchable'
            ]
        },
        {
            'title': 'AI-Enhanced Content Analysis',
            'story': 'As a business analyst, I want AI to understand document context '
                    'so that I can get more accurate feature prioritization.',
            'acceptance': [
                'LLM analyzes document content for semantic meaning',
                'System identifies key themes and priorities',
                'AI provides confidence scores for extracted insights',
                'Fallback processing handles edge cases'
            ]
        },
        {
            'title': 'Framework-Based Prioritization',
            'story': 'As a product owner, I want to apply different prioritization frameworks '
                    '(RICE, ICE, MoSCoW) so that I can align with my team\'s methodology.',
            'acceptance': [
                'User can select from multiple frameworks',
                'System applies framework-specific scoring',
                'Results are presented in framework format',
                'Custom frameworks can be configured'
            ]
        }
    ]
    
    for story in stories:
        doc.add_heading(story['title'], level=2)
        doc.add_paragraph(f"Story: {story['story']}")
        doc.add_paragraph("Acceptance Criteria:")
        for criteria in story['acceptance']:
            doc.add_paragraph(f"• {criteria}")
        doc.add_paragraph("")  # Add spacing
    
    # Technical Specifications
    doc.add_heading('Technical Specifications', level=1)
    
    doc.add_heading('Architecture Components', level=2)
    arch_table = doc.add_table(rows=1, cols=3)
    arch_table.style = 'Table Grid'
    hdr_cells = arch_table.rows[0].cells
    hdr_cells[0].text = 'Component'
    hdr_cells[1].text = 'Technology'
    hdr_cells[2].text = 'Purpose'
    
    components = [
        ['API Layer', 'Django REST Framework', 'Request handling and validation'],
        ['Workflow Engine', 'LangGraph', 'AI pipeline orchestration'],
        ['Document Processor', 'Hybrid LLM + Traditional', 'Content extraction'],
        ['Database', 'PostgreSQL', 'Data persistence'],
        ['Cache Layer', 'Redis', 'Performance optimization']
    ]
    
    for comp in components:
        row_cells = arch_table.add_row().cells
        row_cells[0].text = comp[0]
        row_cells[1].text = comp[1]
        row_cells[2].text = comp[2]
    
    doc.save(filename)
    print(f"Created DOCX document: {filename}")

def create_xlsx_document(filename):
    """Create a comprehensive XLSX document with feature prioritization data"""
    wb = Workbook()
    
    # Feature Backlog Sheet
    ws1 = wb.active
    ws1.title = "Feature Backlog"
    
    # Headers
    headers = ["Feature ID", "Feature Name", "Description", "Priority", "Effort", 
               "Impact", "RICE Score", "Quarter", "Status", "Owner"]
    for col, header in enumerate(headers, 1):
        ws1.cell(row=1, column=col, value=header)
    
    # Sample data
    features = [
        ["F001", "Hybrid LLM Processing", "Combine traditional and AI parsing", "High", 8, 9, 18, "Q1 2024", "In Progress", "AI Team"],
        ["F002", "Real-time Collaboration", "Multi-user document editing", "Medium", 13, 7, 15, "Q2 2024", "Planning", "Frontend Team"],
        ["F003", "Advanced Analytics", "Processing metrics dashboard", "High", 5, 8, 16, "Q1 2024", "Ready", "Data Team"],
        ["F004", "API Rate Limiting", "Prevent abuse and ensure stability", "Low", 3, 5, 8, "Q4 2024", "Backlog", "Backend Team"],
        ["F005", "Multi-format Support", "Support for additional file types", "Medium", 8, 6, 12, "Q3 2024", "Research", "AI Team"],
        ["F006", "User Authentication", "OAuth and SSO integration", "High", 5, 9, 18, "Q1 2024", "Complete", "Security Team"],
        ["F007", "Mobile App", "iOS and Android applications", "Low", 21, 4, 6, "Q1 2025", "Idea", "Mobile Team"],
        ["F008", "Batch Processing", "Handle large document sets", "Medium", 8, 7, 14, "Q2 2024", "Planning", "Backend Team"]
    ]
    
    for row, feature in enumerate(features, 2):
        for col, value in enumerate(feature, 1):
            ws1.cell(row=row, column=col, value=value)
    
    # Technical Debt Sheet
    ws2 = wb.create_sheet("Technical Debt")
    debt_headers = ["Debt ID", "Component", "Issue", "Severity", "Effort", "Impact on Performance"]
    for col, header in enumerate(debt_headers, 1):
        ws2.cell(row=1, column=col, value=header)
    
    debt_items = [
        ["TD001", "Document Parser", "Legacy PDF parsing library", "Medium", 5, "20% slower processing"],
        ["TD002", "Database", "Missing indexes on queries", "High", 2, "3x slower API responses"],
        ["TD003", "API", "No request validation", "High", 3, "Security vulnerability"],
        ["TD004", "Frontend", "Outdated React version", "Low", 8, "Maintenance burden"],
        ["TD005", "Monitoring", "No application metrics", "Medium", 4, "No visibility into issues"]
    ]
    
    for row, debt in enumerate(debt_items, 2):
        for col, value in enumerate(debt, 1):
            ws2.cell(row=row, column=col, value=value)
    
    # Metrics Sheet
    ws3 = wb.create_sheet("Success Metrics")
    metrics_headers = ["Metric", "Current Value", "Target Value", "Timeline", "Measurement Method"]
    for col, header in enumerate(metrics_headers, 1):
        ws3.cell(row=1, column=col, value=header)
    
    metrics = [
        ["Processing Accuracy", "75%", "95%", "Q2 2024", "Manual validation of 100 random documents"],
        ["Processing Speed", "30 sec/doc", "5 sec/doc", "Q1 2024", "Average processing time measurement"],
        ["User Satisfaction", "3.2/5", "4.5/5", "Q3 2024", "Quarterly user survey"],
        ["System Uptime", "95%", "99.9%", "Q2 2024", "Automated monitoring alerts"],
        ["API Response Time", "500ms", "100ms", "Q1 2024", "Performance monitoring tools"]
    ]
    
    for row, metric in enumerate(metrics, 2):
        for col, value in enumerate(metric, 1):
            ws3.cell(row=row, column=col, value=value)
    
    wb.save(filename)
    print(f"Created XLSX document: {filename}")

def main():
    """Create all test documents"""
    base_dir = "/home/moemoeen/Documents/GitHub/Python_Projects_Personal/productroadmap_ai/test_documents"
    
    # Install required packages for PDF creation
    try:
        from reportlab.lib.pagesizes import letter
        print("ReportLab is available for PDF creation")
    except ImportError:
        print("Installing ReportLab for PDF creation...")
        os.system("pip install reportlab")
    
    try:
        # Create test documents
        create_pdf_document(os.path.join(base_dir, "product_roadmap.pdf"))
        create_docx_document(os.path.join(base_dir, "user_stories.docx"))
        create_xlsx_document(os.path.join(base_dir, "feature_backlog.xlsx"))
        
        print("\n✅ All test documents created successfully!")
        print(f"📂 Documents saved in: {base_dir}")
        
        # List created files
        for filename in os.listdir(base_dir):
            if filename.endswith(('.pdf', '.docx', '.xlsx')):
                filepath = os.path.join(base_dir, filename)
                size = os.path.getsize(filepath)
                print(f"   📄 {filename} ({size:,} bytes)")
                
    except Exception as e:
        print(f"❌ Error creating test documents: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()
